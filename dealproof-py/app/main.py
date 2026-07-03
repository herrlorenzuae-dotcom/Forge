"""DealProof — FastAPI app. Project-centric flow:
landing (projects) → new project (name) → upload the KYC document → analysis
report (where each answer can come from) → structure / requests."""
import os
from fastapi import FastAPI, Request, Form, UploadFile, File, Query
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from . import config
from .db import init_db, db, rows, one
from .engine import projects as proj
from .engine import spa as spa_engine
from .engine import transparency
from .engine import profile as profile_engine
from .engine import workflow
from .engine import export as exporter
from .engine.structure import get_structure
from .engine.orgchart import build_orgchart, render_svg, default_subject
from .engine.coverage import build_coverage
from .engine.analysis import build_analysis
from .engine import requests as reqs
from .engine import mapping
from .engine import templatefill
from .engine.brain import brain_stats, get_brain_options, learn_from_document, update_entry, delete_entry

BASE = os.path.dirname(__file__)
app = FastAPI(title="DealProof")
app.mount("/static", StaticFiles(directory=os.path.join(BASE, "static")), name="static")
templates = Jinja2Templates(directory=os.path.join(BASE, "templates"))


# ── Tenant (Mandant) login: every session belongs to one tenant, and each
# tenant works on its OWN database file — projects, structures and the KYC
# Brain never cross tenants. The mock login is designed to be swapped for SSO
# or real user accounts later (only this block changes).
def _tenant_by_slug(slug: str):
    return next((t for t in config.TENANTS if t["slug"] == slug), None)


def _tenant_token(t) -> str:
    import hashlib
    return hashlib.sha256(f"dealproof:{t['slug']}:{t['password']}".encode()).hexdigest()


def _init_tenant(t) -> str:
    """Point this request at the tenant's database, creating it on first use."""
    os.makedirs(config.DATA_DIR, exist_ok=True)
    path = config.tenant_db_path(t["slug"])
    from .db import set_db_path
    set_db_path(path)
    if path not in _initialized_dbs:
        init_db()
        from .demo import seed_demo
        seed_demo()
        _initialized_dbs.add(path)
    return path


_initialized_dbs = set()


@app.middleware("http")
async def _require_tenant(request: Request, call_next):
    path = request.url.path
    if path.startswith("/static") or path in ("/login", "/api/health"):
        return await call_next(request)
    t = _tenant_by_slug(request.cookies.get("dp_tenant", ""))
    if not t or request.cookies.get("dp_auth") != _tenant_token(t):
        return RedirectResponse("/login", status_code=303)
    _init_tenant(t)
    request.state.tenant = t
    return await call_next(request)


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return templates.TemplateResponse(request, "login.html",
        {"request": request, "error": "", "tenants": config.TENANTS})


@app.post("/login")
async def do_login(request: Request, tenant: str = Form(""), password: str = Form("")):
    t = _tenant_by_slug(tenant) or (config.TENANTS[0] if len(config.TENANTS) == 1 else None)
    if t and password == t["password"]:
        resp = RedirectResponse("/", status_code=303)
        resp.set_cookie("dp_tenant", t["slug"], httponly=True, samesite="lax", max_age=60 * 60 * 24 * 30)
        resp.set_cookie("dp_auth", _tenant_token(t), httponly=True, samesite="lax", max_age=60 * 60 * 24 * 30)
        return resp
    return templates.TemplateResponse(request, "login.html",
        {"request": request, "error": "Wrong password.", "tenants": config.TENANTS})


@app.get("/logout")
def logout():
    resp = RedirectResponse("/login", status_code=303)
    resp.delete_cookie("dp_tenant")
    resp.delete_cookie("dp_auth")
    return resp


@app.on_event("startup")
def _startup():
    init_db()
    from .demo import seed_demo
    seed_demo()  # one clearly-labelled demo project (idempotent)


def ctx(request, **kw):
    out = {"request": request, "has_key": config.HAS_KEY,
           "tenant": getattr(request.state, "tenant", None), **kw}
    # The guided step rail renders on EVERY page that has a project in context —
    # injected here centrally so no route can forget it. The highlighted step is
    # the page the user is ON (the Overview hosts steps 1+2: highlight Project
    # until the subject company is set, then Questionnaire).
    p = kw.get("project")
    if p and "wf" not in kw:
        act = kw.get("active", "")
        if act == "project":
            step_key = "project" if not (p.get("subject_company") or "").strip() else "questionnaire"
        else:
            step_key = {"data": "data", "structure": "structure", "deliver": "deliver"}.get(act, "")
        out["wf"] = workflow.steps(p["id"], step_key)
    return out


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        try:
            import fitz
            doc = fitz.open(stream=data, filetype="pdf")
            return "\n".join(p.get_text() for p in doc)
        except Exception:
            return ""
    if name.endswith(".docx"):
        try:
            from io import BytesIO
            from docx import Document
            d = Document(BytesIO(data))
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for r in t.rows:
                    parts.append("\t".join(c.text for c in r.cells))
            return "\n".join(parts)
        except Exception:
            return ""
    try:
        return data.decode("utf-8")
    except Exception:
        return data.decode("latin-1", "ignore")


def one_qn(qid):
    with db() as con:
        return one(con, "SELECT * FROM questionnaires WHERE id=?", (qid,))


def rows_questions(qid):
    with db() as con:
        qs = rows(con, "SELECT * FROM questions WHERE questionnaire_id=? ORDER BY position", (qid,))
        ans = {a["question_id"]: a for a in rows(con, "SELECT a.* FROM answers a JOIN questions q ON q.id=a.question_id WHERE q.questionnaire_id=?", (qid,))}
    for q in qs:
        q["answer"] = ans.get(q["id"])
        q["options"] = get_brain_options(q["prompt"])
    return qs


# ── Landing / projects ──
@app.get("/", response_class=HTMLResponse)
def landing(request: Request):
    from .engine import connectors
    return templates.TemplateResponse(request, "landing.html", ctx(request, active="projects",
        projects=proj.list_projects(), connectors=connectors.status(), model=config.MODEL))


@app.post("/projects")
def create_project(name: str = Form(...), subject_company: str = Form(""),
                   register_no: str = Form(""), portfolio_company: str = Form("")):
    pid = proj.create_project(name, subject_company, register_no, portfolio_company)
    return RedirectResponse(f"/projects/{pid}", status_code=303)


@app.post("/projects/{pid}/edit")
def edit_project(pid: str, name: str = Form(""), subject_company: str = Form(""),
                 register_no: str = Form(""), portfolio_company: str = Form("")):
    proj.update_project(pid, name=name, subject_company=subject_company,
                        register_no=register_no, portfolio_company=portfolio_company)
    return RedirectResponse(f"/projects/{pid}", status_code=303)


@app.post("/projects/{pid}/delete")
def delete_project(pid: str):
    proj.delete_project(pid)
    return RedirectResponse("/", status_code=303)


# ── Entity profile (master-data sheet) ──
@app.get("/projects/{pid}/profile", response_class=HTMLResponse)
def profile_page(request: Request, pid: str, pulled: int = Query(None)):
    return templates.TemplateResponse(request, "profile.html", ctx(request,
        active="data", project=proj.get_project(pid),
        profile=profile_engine.get_profile(pid), pulled=pulled, wf=workflow.steps(pid, "data")))


@app.post("/projects/{pid}/profile")
async def profile_save(request: Request, pid: str):
    form = await request.form()
    profile_engine.save_profile(pid, {k: v for k, v in form.items()})
    proj.touch(pid)
    return RedirectResponse(f"/projects/{pid}/profile", status_code=303)


@app.post("/projects/{pid}/profile/pull")
def profile_pull(pid: str):
    res = profile_engine.pull_profile(pid)
    proj.touch(pid)
    return RedirectResponse(f"/projects/{pid}/profile?pulled={res['pulled']}", status_code=303)


@app.get("/projects/{pid}", response_class=HTMLResponse)
def project_page(request: Request, pid: str):
    project = proj.get_project(pid)
    if not project:
        return RedirectResponse("/")
    return templates.TemplateResponse(request, "project.html",
        ctx(request, active="project", project=project, documents=proj.list_documents(pid)))


@app.post("/projects/{pid}/documents")
async def upload_document(pid: str, file: UploadFile = File(None), pasted: str = Form(""), requester: str = Form("")):
    text, filename, content = pasted, "pasted.txt", b""
    if file is not None and file.filename:
        content = await file.read()
        text = extract_text(file.filename, content)
        filename = file.filename
    if not text.strip():
        return RedirectResponse(f"/projects/{pid}", status_code=303)
    res = proj.add_document(pid, filename, text, requester, content=content)
    return RedirectResponse(f"/projects/{pid}/analysis/{res['questionnaire_id']}", status_code=303)


# ── Analysis ──
@app.get("/projects/{pid}/analysis/{qid}", response_class=HTMLResponse)
def analysis_page(request: Request, pid: str, qid: str):
    analysis = build_analysis(qid)
    questions = rows_questions(qid)
    # one unified list: plan (where the answer can come from) + actual state
    plan = {it["questionId"]: it for it in analysis["items"]}
    req_by_q = {r["question_id"]: r for r in reqs.list_requests(pid)}
    for q in questions:
        q["plan"] = plan.get(q["id"], {})
        q["req"] = req_by_q.get(q["id"])
        if q["answer"] and (q["answer"]["value"] or "").strip():
            q["state"] = "answered"
        elif q["req"] or q["plan"].get("source") == "request":
            q["state"] = "manual"
        else:
            q["state"] = "open"
    return templates.TemplateResponse(request, "analysis.html", ctx(request,
        active="project", project=proj.get_project(pid), qn=one_qn(qid), questions=questions,
        analysis=analysis, coverage=build_coverage(qid), fillable=templatefill.is_fillable(qid),
        requests=reqs.list_requests(pid), request_text=reqs.render_request_list(pid),
        wf=workflow.steps(pid, "answers")))


@app.get("/projects/{pid}/analysis/{qid}/export.docx")
def export_qn_docx(pid: str, qid: str):
    return Response(exporter.questionnaire_docx(qid),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="questionnaire-{qid}.docx"'})


@app.get("/projects/{pid}/analysis/{qid}/export.pdf")
def export_qn_pdf(pid: str, qid: str):
    return Response(exporter.questionnaire_pdf(qid), media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="questionnaire-{qid}.pdf"'})


@app.get("/projects/{pid}/analysis/{qid}/export.xlsx")
def export_qn_xlsx(pid: str, qid: str):
    return Response(exporter.questionnaire_xlsx(qid),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="questionnaire-{qid}.xlsx"'})


@app.get("/projects/{pid}/analysis/{qid}/fill")
def fill_original_doc(pid: str, qid: str):
    res = templatefill.fill_original(qid)
    if not res:
        return RedirectResponse(f"/projects/{pid}/analysis/{qid}", status_code=303)
    body, mt, fname = res
    return Response(body, media_type=mt, headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@app.post("/projects/{pid}/analysis/{qid}/answer")
def do_answer_all(pid: str, qid: str):
    mapping.answer_all(qid)
    return RedirectResponse(f"/projects/{pid}/analysis/{qid}", status_code=303)


@app.post("/projects/{pid}/analysis/{qid}/requests")
def do_requests(pid: str, qid: str):
    reqs.generate_requests(qid)
    return RedirectResponse(f"/projects/{pid}/analysis/{qid}", status_code=303)


@app.post("/projects/{pid}/analysis/{qid}/finalize")
def do_finalize(pid: str, qid: str):
    mapping.finalize(qid)
    return RedirectResponse(f"/projects/{pid}/analysis/{qid}", status_code=303)


@app.post("/questions/{question_id}/answer")
def do_answer_one(question_id: str, pid: str = Form(...), qid: str = Form(...)):
    mapping.answer_question(question_id)
    return RedirectResponse(f"/projects/{pid}/analysis/{qid}", status_code=303)


@app.post("/questions/{question_id}/set")
def do_set(question_id: str, value: str = Form(""), pid: str = Form(...), qid: str = Form(...)):
    mapping.set_answer(question_id, value)
    return RedirectResponse(f"/projects/{pid}/analysis/{qid}", status_code=303)


@app.post("/requests/{req_id}")
def do_update_request(req_id: str, status: str = Form(...), pid: str = Form(...), qid: str = Form(...)):
    reqs.update_request(req_id, status=status)
    return RedirectResponse(f"/projects/{pid}/analysis/{qid}", status_code=303)


# ── Structure ──
@app.get("/projects/{pid}/structure", response_class=HTMLResponse)
def structure_page(request: Request, pid: str, view: str = Query("excerpt"), subject: str = Query(None)):
    chart = build_orgchart(pid)
    excerpt = view == "excerpt"
    subj = subject or default_subject(chart["nodes"], chart["edges"])
    other_projects = [p for p in proj.list_projects() if p["id"] != pid]
    from .engine.derivation import ubo_derivation
    return templates.TemplateResponse(request, "structure.html", ctx(request,
        active="structure", project=proj.get_project(pid), structure=get_structure(pid),
        chart=chart, chart_svg=render_svg(pid, subject=subj, excerpt=excerpt),
        view=view, subject=subj, spec=spa_engine.structure_to_spec(pid),
        derivation=ubo_derivation(pid, subj),
        other_projects=other_projects, wf=workflow.steps(pid, "structure")))


@app.post("/projects/{pid}/structure")
async def ingest_structure(pid: str, file: UploadFile = File(None), pasted: str = Form(""),
                           replace: str = Form(""), mode: str = Form("replace"),
                           attach_to: str = Form(""), attach_rel: str = Form("")):
    from .engine import chartpdf, pptxchart
    text, spec = pasted, None
    if file is not None and file.filename:
        data = await file.read()
        # drawn charts are read geometrically, fully local: PowerPoint decks and
        # vector PDF charts
        if file.filename.lower().endswith(".pptx"):
            spec = pptxchart.extract_spec(data)
        elif file.filename.lower().endswith(".pdf") and chartpdf.is_chart_pdf(data):
            spec = chartpdf.extract_spec(data)
        if spec is None:
            text = extract_text(file.filename, data)
    if spec is None and text.strip():
        # An edited spec (manual correction) is the deterministic format — parse it
        # directly; an uploaded SPA goes through the (model-assisted) extractor.
        if replace == "spec":
            spec = spa_engine.parse_structure_spec(text)
        else:
            spec = spa_engine.extract_from_spa(text, pid)
    if spec:
        if mode == "merge" and replace != "spec":
            spa_engine.merge_structure(pid, spec, attach_to=attach_to, attach_rel=attach_rel)
        else:
            spa_engine.apply_structure(pid, spec)
        proj.touch(pid)
    return RedirectResponse(f"/projects/{pid}/structure?view={'full' if replace=='spec' or mode=='merge' else 'excerpt'}", status_code=303)


@app.post("/projects/{pid}/structure/copy")
def copy_structure(pid: str, source_project: str = Form(...)):
    if source_project and source_project != pid:
        spa_engine.copy_structure(pid, source_project)
        proj.touch(pid)
    return RedirectResponse(f"/projects/{pid}/structure?view=full", status_code=303)


@app.post("/projects/{pid}/ubos")
async def import_ubos(pid: str, file: UploadFile = File(None), pasted: str = Form("")):
    text = pasted
    if file is not None and file.filename:
        data = await file.read()
        text = extract_text(file.filename, data)
    if text.strip():
        transparency.import_extract(pid, text)
        proj.touch(pid)
    return RedirectResponse(f"/projects/{pid}/structure?view=excerpt", status_code=303)


# ── Structure / chart export ──
@app.get("/projects/{pid}/structure.svg")
def export_svg(pid: str, view: str = Query("full")):
    body = exporter.chart_svg(pid, excerpt=view == "excerpt")
    return Response(body, media_type="image/svg+xml",
                    headers={"Content-Disposition": f'attachment; filename="structure-{pid}.svg"'})


@app.get("/projects/{pid}/structure.png")
def export_png(pid: str, view: str = Query("full")):
    body = exporter.chart_png(pid, excerpt=view == "excerpt")
    return Response(body, media_type="image/png",
                    headers={"Content-Disposition": f'attachment; filename="structure-{pid}.png"'})


@app.get("/projects/{pid}/structure.pdf")
def export_structure_pdf(pid: str, view: str = Query("full")):
    body = exporter.chart_pdf(pid, excerpt=view == "excerpt")
    return Response(body, media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="structure-{pid}.pdf"'})


@app.get("/projects/{pid}/structure.xlsx")
def export_structure_xlsx(pid: str):
    body = exporter.structure_xlsx(pid)
    return Response(body, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": f'attachment; filename="structure-{pid}.xlsx"'})


# ── Deliver (final step: filled original + chart files) ──
@app.get("/projects/{pid}/deliver", response_class=HTMLResponse)
def deliver_page(request: Request, pid: str, qn_id: str = Query(None)):
    with db() as con:
        qns = rows(con, "SELECT * FROM questionnaires WHERE client_id=? ORDER BY created_at DESC", (pid,))
        qn = next((q for q in qns if q["id"] == qn_id), qns[0] if qns else None)
        stats = {"answered": 0, "total": 0, "manual": 0, "entities": 0}
        if qn:
            stats["total"] = one(con, "SELECT COUNT(*) c FROM questions WHERE questionnaire_id=?", (qn["id"],))["c"]
            stats["answered"] = one(con, """SELECT COUNT(*) c FROM answers a JOIN questions q ON q.id=a.question_id
                                            WHERE q.questionnaire_id=? AND a.value!=''""", (qn["id"],))["c"]
            stats["manual"] = one(con, "SELECT COUNT(*) c FROM info_requests WHERE questionnaire_id=?", (qn["id"],))["c"]
        stats["entities"] = one(con, "SELECT COUNT(*) c FROM entities WHERE client_id=?", (pid,))["c"]
    return templates.TemplateResponse(request, "deliver.html", ctx(request,
        active="deliver", project=proj.get_project(pid), qn=qn, qns=qns, stats=stats,
        fillable=templatefill.is_fillable(qn["id"]) if qn else False,
        wf=workflow.steps(pid, "deliver")))


@app.post("/projects/{pid}/questionnaires/{qid}/review")
def mark_reviewed(pid: str, qid: str, reviewer: str = Form(...)):
    with db() as con:
        con.execute("UPDATE questionnaires SET reviewed_by=?, reviewed_at=date('now') WHERE id=?",
                    (reviewer.strip(), qid))
    proj.touch(pid)
    return RedirectResponse(f"/projects/{pid}/deliver?qn_id={qid}", status_code=303)


# ── Backup (per tenant — the file IS the tenant's complete data) ──
@app.get("/backup")
def backup(request: Request):
    from datetime import date
    from .db import current_db_path
    slug = getattr(request.state, "tenant", {}).get("slug", "workspace")
    data = open(current_db_path(), "rb").read()
    return Response(data, media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="dealproof-{slug}-{date.today().isoformat()}.db"'})


# ── Brain (cross-project memory) ──
@app.get("/brain", response_class=HTMLResponse)
def brain_page(request: Request, learned: int = Query(None), total: int = Query(None),
               kind: str = Query(""), company: str = Query("")):
    return templates.TemplateResponse(request, "brain.html", ctx(request, active="brain",
        entries=brain_stats(), learned=learned, total=total, kind=kind, company=company))


@app.post("/brain/learn")
async def brain_learn(file: UploadFile = File(None)):
    if file is not None and file.filename:
        data = await file.read()
        res = learn_from_document(file.filename, data, extract_text(file.filename, data))
        import urllib.parse
        extra = f"&kind={res.get('kind','')}&company={urllib.parse.quote(res.get('company',''))}" if res.get("kind") else ""
        return RedirectResponse(f"/brain?learned={res['learned']}&total={res['total']}{extra}", status_code=303)
    return RedirectResponse("/brain", status_code=303)


@app.post("/brain/{entry_id}/update")
def brain_update(entry_id: str, value: str = Form(...)):
    update_entry(entry_id, value)
    return RedirectResponse("/brain", status_code=303)


@app.post("/brain/{entry_id}/delete")
def brain_delete(entry_id: str):
    delete_entry(entry_id)
    return RedirectResponse("/brain", status_code=303)


# ── JSON API ──
@app.get("/api/health")
def health():
    from .engine import connectors
    return {"ok": True,
            "anthropic": {"enabled": config.HAS_KEY, "model": config.MODEL},
            "connectors": connectors.status()}


@app.get("/api/projects/{pid}/analysis/{qid}")
def api_analysis(pid: str, qid: str):
    return build_analysis(qid)


# Lets you start the app with plain `python -m app.main` (or `python app/main.py`).
# Port via KYC_PORT env var, default 8000.
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app.main:app", host="127.0.0.1", port=int(os.environ.get("KYC_PORT", "3100")), reload=True)
