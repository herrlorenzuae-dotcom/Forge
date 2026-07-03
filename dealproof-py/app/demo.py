"""Optional demo project with dummy data — created once on startup so a new
user can see the full flow (structure incl. Komplementär control, a parsed
questionnaire, and an analysis spanning every source). It is just one entry in
the project list; delete it and it only returns on a fresh database."""
from .db import db, one, gen_id
from .engine import spa, projects, brain

DEMO_ID = "proj_demo"
DEMO_NAME = "Demo — Project Cedar (acquisition of Helio Thermal Systems)"

# A realistic, FICTIONAL PE structure: a GmbH &
# Co. KG cascade where each KG is controlled by its general partner (Komplementär)
# with no equity, running from the portfolio company up through the fund and
# feeder vehicles to the individual partners. Names/figures are invented.
SPEC = """\
# Asset and acquisition vehicle
Cedar BidCo S.à r.l. -> Helio Thermal Systems GmbH : 94%
# Deal fund (the project vehicle), controlled by its Komplementär
Cedar Blue GP GmbH -> Cedar Blue GmbH & Co. geschlossene Investment KG : control (General partner / Komplementär)
Cedar Blue GmbH & Co. geschlossene Investment KG -> Cedar BidCo S.à r.l. : 100%
Cedar Blue Initiators GmbH & Co. KG -> Cedar Blue GmbH & Co. geschlossene Investment KG : 8%
Cedar Blue Team GmbH & Co. KG -> Cedar Blue GmbH & Co. geschlossene Investment KG : 5%
Cedar Beteiligungen GmbH & Co. KG -> Cedar Blue GmbH & Co. geschlossene Investment KG : 87%
# Main fund holding, controlled by its Komplementär
Cedar Beteiligungen Verwaltungs GmbH -> Cedar Beteiligungen GmbH & Co. KG : control (General partner / Komplementär)
Cedar GmbH & Co. KG -> Cedar Beteiligungen GmbH & Co. KG : 100%
# Top partnership, controlled by its Komplementär
Cedar Verwaltungs GmbH -> Cedar GmbH & Co. KG : control (General partner / Komplementär)
# Partners (limited partners) hold the economic interest in the top KG
Dr. Anna Vogt -> Cedar GmbH & Co. KG : 28%
Maximilian Stein -> Cedar GmbH & Co. KG : 22%
Sofia Brandt -> Cedar GmbH & Co. KG : 18%
Further partners (11) -> Cedar GmbH & Co. KG : 32%
# Control of the group runs through the managing partners of the top Komplementär
Dr. Anna Vogt -> Cedar Verwaltungs GmbH : control (Managing partner)
Maximilian Stein -> Cedar Verwaltungs GmbH : control (Managing partner)
TARGET: Helio Thermal Systems GmbH
UBO: Dr. Anna Vogt
UBO: Maximilian Stein
"""

# (entity name, key, value, source) — a couple of facts already "on file"
ATTRS = [
    ("Cedar BidCo S.à r.l.", "LEI", "391200CEDARBIDCO0007", "GLEIF"),
]

# prior verified answers folded into the Brain (so some questions resolve there)
BRAIN = [
    ("Full legal name of the contracting entity?", "Cedar BidCo S.à r.l.", 3),
    ("Is the contracting entity a regulated financial institution?", "No — special-purpose acquisition vehicle, not a supervised institution.", 2),
    ("What is the source of funds for the transaction?",
     "Equity drawn from Cedar Blue GmbH & Co. geschlossene Investment KG (partner commitments) plus senior acquisition financing.", 2),
]

# A deliberately MESSY questionnaire, the way a real bank PDF/Word export reads:
# section headings, questions wrapped across lines, numbering in its own column,
# and trailing "Y / N" answer cells. The intake reflow turns this back into clean
# atomic questions grouped under their sections.
QUESTIONNAIRE = """\
KYC & ONBOARDING QUESTIONNAIRE
Banque de Genève SA — Correspondent Onboarding

SECTION 1 — ENTITY & OWNERSHIP

1.1   Full legal name of the
      contracting entity?
1.2   Legal Entity Identifier
      (LEI)?
1.3   Registered office
      address?
1.4   Date of incorporation?
1.5   Is the contracting entity a regulated
      financial institution?                          Y / N
1.6   Identify all ultimate beneficial owners
      holding 25% or more.

SECTION 2 — SCREENING & RISK

2.1   Is any beneficial owner a politically
      exposed person (PEP)?                            Y / N
2.2   Tax residence of the contracting
      entity?

SECTION 3 — TRANSACTION & RELATIONSHIP

3.1   What is the source of funds for the
      transaction?
3.2   What is the purpose of the business
      relationship?

SECTION 4 — DOCUMENTATION

4.1   Please attach a certified copy of the
      passport for each UBO.
"""


def _demo_word(text: str) -> bytes:
    """A fillable Word version of the questionnaire (clean question paragraphs,
    each followed by a blank answer line) so the demo can show 'Fill original'."""
    from io import BytesIO
    from docx import Document
    from .engine.intake import parse_questions
    doc = Document()
    doc.add_heading("KYC & Onboarding Questionnaire", 0)
    doc.add_paragraph("Banque de Genève SA — Correspondent Onboarding")
    cur = None
    for q in parse_questions(text):
        if q["section"] != cur:
            cur = q["section"]
            doc.add_heading(cur.title(), level=1)
        doc.add_paragraph(q["prompt"])
        doc.add_paragraph("")  # answer line — filled in by "Fill original"
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# UBO legal basis for the demo, as a real Transparenzregister extract would
# state it: (entity name, Art des wirtschaftlichen Interesses, Umfang)
UBO_BASIS = {
    "Dr. Anna Vogt": ("Beteiligung an der Vereinigung, Höhe der Kapitalanteile (§ 19 Abs. 3 Nr. 1a GwG)", "28 %"),
    "Maximilian Stein": ("Ausübung von Kontrolle auf sonstige Weise als Managing Partner der Cedar Verwaltungs GmbH", ""),
}


def seed_demo() -> None:
    with db() as con:
        existing = one(con, "SELECT 1 FROM clients WHERE id=?", (DEMO_ID,))
        current = one(con, "SELECT 1 FROM ubos WHERE client_id=? AND note!='' LIMIT 1", (DEMO_ID,)) if existing else None
    if existing and current:
        return
    if existing:
        # seeded by an OLDER version (no UBO legal-basis notes) → refresh once,
        # so pulled updates are actually visible in the demo
        projects.delete_project(DEMO_ID)
    with db() as con:
        con.execute("""INSERT INTO clients (id, name, subject_company, register_no, portfolio_company, status, updated_at)
                       VALUES (?,?,?,?,?, 'open', datetime('now'))""",
                    (DEMO_ID, DEMO_NAME, "Cedar BidCo S.à r.l.", "RCS Luxembourg B 271 904",
                     "Helio Thermal Systems GmbH"))

    spa.apply_structure(DEMO_ID, spa.parse_structure_spec(SPEC))

    with db() as con:
        for ent_name, key, value, src in ATTRS:
            row = one(con, "SELECT id FROM entities WHERE client_id=? AND name=?", (DEMO_ID, ent_name))
            if row:
                con.execute("INSERT INTO entity_attributes (id, entity_id, key, value, source, as_of) VALUES (?,?,?,?,?, date('now'))",
                            (gen_id("attr"), row["id"], key, value, src))
        # UBO records carry the legal basis (§ note), exactly like a register import
        from .engine.ubolaw import classify_extent
        for ent_name, (art, umfang) in UBO_BASIS.items():
            row = one(con, "SELECT id FROM entities WHERE client_id=? AND name=?", (DEMO_ID, ent_name))
            if row:
                c = classify_extent(art, umfang)
                con.execute("UPDATE ubos SET basis=?, pct=?, note=? WHERE client_id=? AND entity_id=?",
                            (c["basis"], c["pct"], c["note"], DEMO_ID, row["id"]))

    for prompt, value, n in BRAIN:
        for _ in range(n):
            brain.record_finalized_answer(prompt, value)

    projects.add_document(DEMO_ID, "Bank-onboarding-KYC.docx", QUESTIONNAIRE,
                          "Banque de Genève SA", content=_demo_word(QUESTIONNAIRE))
