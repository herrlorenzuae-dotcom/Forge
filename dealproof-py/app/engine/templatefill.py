"""Fill the answers back into the EXACT original document the bank sent —
preserving its layout — rather than generating a fresh one.

  - PDF AcroForm: match each form field (by its name and the label printed
    next to it) to an answered question and write the value into the widget.
  - Word (.docx): fill {{placeholders}}, answer cells in a Question/Answer
    table, and write the answer beneath a matching question paragraph.

Matching is fuzzy (token overlap of the question against the field label), so
it works on real forms whose field names don't exactly equal our prompts.
Returns (bytes, media_type, filename) or None when the original can't be
filled in place (e.g. a flat PDF/scan or a plain-text upload) — the caller
falls back to the generated export."""
import re
from io import BytesIO
from ..db import db, rows, one

STOP = {"the", "of", "a", "an", "is", "are", "to", "for", "and", "or", "in",
        "on", "any", "all", "each", "please", "provide", "entity", "contracting",
        "this", "that", "with", "as", "be", "your", "you", "do", "does", "has",
        "have", "what", "which", "who", "name"}


def _norm(s: str) -> str:
    return re.sub(r"[^a-z0-9 ]", " ", (s or "").lower())


def _tokens(s: str) -> set:
    return {w for w in _norm(s).split() if len(w) > 2 and w not in STOP}


def _answered(questionnaire_id: str):
    """[(prompt, value, kind, token_set)] for questions that have an answer."""
    with db() as con:
        qs = rows(con, "SELECT * FROM questions WHERE questionnaire_id=? ORDER BY position", (questionnaire_id,))
        ans = {a["question_id"]: a for a in rows(con,
               "SELECT a.* FROM answers a JOIN questions q ON q.id=a.question_id WHERE q.questionnaire_id=?",
               (questionnaire_id,))}
    out = []
    for q in qs:
        a = ans.get(q["id"])
        if a and (a["value"] or "").strip():
            out.append((q["prompt"], a["value"], q["kind"], _tokens(q["prompt"])))
    return out


def _best(label: str, answered):
    """The answered question whose tokens are best covered by the label."""
    lt = _tokens(label)
    if not lt:
        return None
    best, score = None, 0.0
    for prompt, val, kind, pt in answered:
        if not pt:
            continue
        cov = len(lt & pt) / len(pt)        # how much of the question the label covers
        cov2 = len(lt & pt) / len(lt)       # …and vice-versa, to reward tight labels
        s = (cov + cov2) / 2
        if s > score:
            best, score = (prompt, val, kind), s
    return best if score >= 0.45 else None


# ── PDF AcroForm ──
def _pdf_label(page, widget) -> str:
    """The field's own name plus any text printed just left of / above it."""
    name = (widget.field_name or "") + " " + (getattr(widget, "field_label", "") or "")
    r = widget.rect
    near = []
    for x0, y0, x1, y1, word, *_ in page.get_text("words"):
        same_row = abs((y0 + y1) / 2 - (r.y0 + r.y1) / 2) < 8 and x1 <= r.x0 + 2
        above = 0 < (r.y0 - y1) < 22 and not (x1 < r.x0 - 120 or x0 > r.x1 + 120)
        if same_row or above:
            near.append((y0, x0, word))
    near.sort()
    return name + " " + " ".join(w for _, _, w in near)


def fill_pdf(content: bytes, answered) -> bytes | None:
    import fitz
    doc = fitz.open(stream=content, filetype="pdf")
    filled = 0
    for page in doc:
        for w in (page.widgets() or []):
            m = _best(_pdf_label(page, w), answered)
            if not m:
                continue
            _, val, kind = m
            try:
                if w.field_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                    yes = bool(re.match(r"\s*(y|yes|true)\b", val, re.I))
                    w.field_value = bool(yes)
                else:
                    w.field_value = str(val)
                w.update()
                filled += 1
            except Exception:
                continue
    if not filled:
        return None
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ── Word (.docx) ──
def _docx_set(paragraph, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def fill_docx(content: bytes, answered) -> bytes | None:
    from docx import Document
    doc = Document(BytesIO(content))
    filled = 0

    # 1) {{placeholder}} tokens anywhere in the body
    ph = re.compile(r"\{\{\s*(.+?)\s*\}\}")
    for p in doc.paragraphs:
        for ph_m in ph.finditer(p.text):
            m = _best(ph_m.group(1), answered)
            if m:
                _docx_set(p, ph.sub(lambda _: m[1], p.text))
                filled += 1
                break

    # 2) Question/Answer tables: first cell = question, fill an empty later cell
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            m = _best(cells[0].text, answered)
            if not m:
                continue
            target = next((c for c in cells[1:] if not c.text.strip()), cells[-1])
            if not target.text.strip():
                _docx_set(target.paragraphs[0], m[1])
                filled += 1

    # 3) Question paragraphs: write the answer into the next empty paragraph
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        t = p.text.strip()
        if not t or not (t.endswith("?") or len(_tokens(t)) >= 2):
            continue
        m = _best(t, answered)
        if not m:
            continue
        nxt = paras[i + 1] if i + 1 < len(paras) else None
        if nxt is not None and not nxt.text.strip():
            _docx_set(nxt, m[1])
            filled += 1

    if not filled:
        return None
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


PDF_MT = "application/pdf"
DOCX_MT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def fill_original(questionnaire_id: str):
    """Return (bytes, media_type, filename) for the filled original, or None."""
    orig = one_original(questionnaire_id)
    if not orig or not orig.get("content"):
        return None
    name = (orig["filename"] or "").lower()
    answered = _answered(questionnaire_id)
    if not answered:
        return None
    base = re.sub(r"\.(pdf|docx)$", "", orig["filename"], flags=re.I)
    if name.endswith(".pdf"):
        b = fill_pdf(orig["content"], answered)
        return (b, PDF_MT, f"{base}-filled.pdf") if b else None
    if name.endswith(".docx"):
        b = fill_docx(orig["content"], answered)
        return (b, DOCX_MT, f"{base}-filled.docx") if b else None
    return None


def one_original(questionnaire_id: str):
    with db() as con:
        return one(con, "SELECT filename, content FROM documents WHERE questionnaire_id=? ORDER BY uploaded_at DESC LIMIT 1",
                   (questionnaire_id,))


def is_fillable(questionnaire_id: str) -> bool:
    orig = one_original(questionnaire_id)
    if not orig or not orig.get("content"):
        return False
    return (orig["filename"] or "").lower().endswith((".pdf", ".docx"))
