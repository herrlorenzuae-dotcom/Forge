"""Export the org chart / structure: SVG, PNG, PDF (rendered from the SVG via
PyMuPDF) and Excel (structure tables, for review/correction).

Also exports the COMPLETED questionnaire — the answers (from Quantium /
YSolutions / the KYC Brain / web / on file) carried back into a document
(Word / PDF / Excel) so the information lands in the questionnaire itself."""
from io import BytesIO
from .orgchart import render_svg
from .structure import get_structure
from ..db import db, rows, one

# answered_by -> how the answer was obtained (shown in the completed document)
SOURCE_LABEL = {
    "source": "From the source document", "brain": "KYC Brain (previously verified)",
    "on_file": "On file", "quantium": "Quantium", "ysolutions": "YSolutions",
    "web": "Web research", "model": "Drafted by model", "human": "Entered by reviewer",
}


def chart_svg(project_id: str, excerpt: bool = False) -> str:
    inner = render_svg(project_id, excerpt=excerpt)
    if not inner.lstrip().startswith("<svg"):
        inner = f'<svg xmlns="http://www.w3.org/2000/svg" width="400" height="60"><text x="8" y="32">No structure</text></svg>'
    return '<?xml version="1.0" encoding="UTF-8"?>\n' + inner


def chart_png(project_id: str, excerpt: bool = False) -> bytes:
    import fitz
    doc = fitz.open(stream=chart_svg(project_id, excerpt).encode("utf-8"), filetype="svg")
    return doc[0].get_pixmap(matrix=fitz.Matrix(2, 2)).tobytes("png")


def chart_pdf(project_id: str, excerpt: bool = False) -> bytes:
    import fitz
    doc = fitz.open(stream=chart_svg(project_id, excerpt).encode("utf-8"), filetype="svg")
    return doc.convert_to_pdf()


def structure_xlsx(project_id: str) -> bytes:
    from openpyxl import Workbook
    s = get_structure(project_id)
    name = {e["id"]: e["name"] for e in s["entities"]}
    wb = Workbook()
    e = wb.active
    e.title = "Entities"
    e.append(["Name", "Kind", "Role", "Jurisdiction", "Registration no.", "Incorporation"])
    for x in s["entities"]:
        e.append([x["name"], x["kind"], x["role"], x["jurisdiction"], x["registration_no"], x["incorporation_date"]])
    r = wb.create_sheet("Relationships")
    r.append(["Parent", "Child", "Percent", "Kind", "Mechanism"])
    for x in s["edges"]:
        r.append([name.get(x["parent_id"], "?"), name.get(x["child_id"], "?"), x["pct"], x["kind"], x["mechanism"]])
    u = wb.create_sheet("UBOs")
    u.append(["Entity", "Basis", "Percent", "PEP", "Residence"])
    for x in s["ubos"]:
        u.append([x["entity_name"], x["basis"], x["pct"], "yes" if x["pep"] else "no", x["residence"]])
    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


# ── Completed questionnaire (answers carried back into the document) ──
def _completed(questionnaire_id: str) -> dict:
    """Questionnaire title + questions (with their answers) grouped by section."""
    with db() as con:
        qn = one(con, "SELECT * FROM questionnaires WHERE id=?", (questionnaire_id,))
        qs = rows(con, "SELECT * FROM questions WHERE questionnaire_id=? ORDER BY position", (questionnaire_id,))
        ans = {a["question_id"]: a for a in rows(con,
               "SELECT a.* FROM answers a JOIN questions q ON q.id=a.question_id WHERE q.questionnaire_id=?",
               (questionnaire_id,))}
    items, answered = [], 0
    for q in qs:
        a = ans.get(q["id"])
        val = (a["value"] if a else "") or ""
        if val:
            answered += 1
        items.append({"section": q["section"] or "General", "prompt": q["prompt"],
                      "answer": val, "source": SOURCE_LABEL.get(a["answered_by"], "") if a and val else "",
                      "status": (a["status"] if a else "") or ""})
    return {"title": (qn["title"] if qn else "Questionnaire"), "items": items,
            "answered": answered, "total": len(items),
            "reviewed_by": (qn["reviewed_by"] if qn else "") or "",
            "reviewed_at": (qn["reviewed_at"] if qn else "") or ""}


def _status_line(d: dict) -> str:
    line = f"Completed questionnaire — {d['answered']} of {d['total']} questions answered."
    if d["reviewed_by"]:
        line += f" Reviewed by {d['reviewed_by']} on {d['reviewed_at']}."
    else:
        line += " Not yet reviewed."
    return line


def questionnaire_docx(questionnaire_id: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    d = _completed(questionnaire_id)
    doc = Document()
    doc.add_heading(d["title"], level=0)
    p = doc.add_paragraph(_status_line(d))
    p.runs[0].italic = True
    cur = None
    for it in d["items"]:
        if it["section"] != cur:
            cur = it["section"]
            doc.add_heading(cur, level=1)
        q = doc.add_paragraph()
        rq = q.add_run(it["prompt"])
        rq.bold = True
        a = doc.add_paragraph()
        ra = a.add_run(it["answer"] or "— to be added manually —")
        if not it["answer"]:
            ra.italic = True
            ra.font.color.rgb = RGBColor(0xB0, 0x50, 0x50)
        if it["source"]:
            s = doc.add_paragraph()
            rs = s.add_run(f"Source: {it['source']}")
            rs.italic = True
            rs.font.size = Pt(8)
            rs.font.color.rgb = RGBColor(0x6D, 0x6A, 0x63)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def questionnaire_xlsx(questionnaire_id: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    d = _completed(questionnaire_id)
    wb = Workbook()
    ws = wb.active
    ws.title = "Questionnaire"
    ws.append([_status_line(d), "", "", "", ""])
    ws.append(["Section", "Question", "Answer", "Source", "Status"])
    for c in ws[2]:
        c.font = Font(bold=True)
    for it in d["items"]:
        ws.append([it["section"], it["prompt"], it["answer"], it["source"], it["status"]])
    widths = [22, 60, 60, 26, 12]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + i)].width = w
    for row in ws.iter_rows(min_row=2):
        for c in row:
            c.alignment = Alignment(wrap_text=True, vertical="top")
    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


def questionnaire_pdf(questionnaire_id: str) -> bytes:
    import fitz
    d = _completed(questionnaire_id)
    doc = fitz.open()
    page = doc.new_page()
    margin, width = 56, 595 - 112
    y = margin

    def line(text, size=10, color=(0.08, 0.2, 0.3), bold=False, gap=4, indent=0):
        nonlocal y, page
        font = "hebo" if bold else "helv"
        # crude word-wrap
        words, cur = text.split(), ""
        lines = []
        for w in words:
            t = (cur + " " + w).strip()
            if fitz.get_text_length(t, fontname=font, fontsize=size) > width - indent:
                lines.append(cur); cur = w
            else:
                cur = t
        if cur:
            lines.append(cur)
        for ln in lines or [""]:
            if y > 780:
                page = doc.new_page(); y = margin
            page.insert_text((margin + indent, y), ln, fontname=font, fontsize=size, color=color)
            y += size + 3
        y += gap

    line(d["title"], size=18, bold=True, gap=2)
    line(_status_line(d), size=9, color=(0.43, 0.42, 0.39), gap=10)
    cur = None
    for it in d["items"]:
        if it["section"] != cur:
            cur = it["section"]
            line(cur.upper(), size=11, bold=True, color=(0, 0.37, 0.72), gap=5)
        line(it["prompt"], size=10, bold=True, gap=1)
        if it["answer"]:
            line(it["answer"], size=10, color=(0.1, 0.1, 0.1), indent=10, gap=1)
            if it["source"]:
                line(f"Source: {it['source']}", size=8, color=(0.43, 0.42, 0.39), indent=10, gap=7)
            else:
                y += 5
        else:
            line("— to be added manually —", size=9, color=(0.69, 0.31, 0.31), indent=10, gap=7)
    return doc.tobytes()
