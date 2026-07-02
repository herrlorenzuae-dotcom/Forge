"""Questionnaire intake — turn a pasted/extracted KYC document into atomic
questions.

Real KYC questionnaires are messy: long questions wrap across several lines when
a PDF is extracted, section headings sit between them, numbering lives in its own
column, and answer cells ("Y / N") trail the text. A naive "every line that ends
in ? is a question" splits one question into fragments and treats headings as
questions.

Two paths:
  1. Model-assisted (when an ANTHROPIC_API_KEY is set): the model reconstructs
     atomic questions, attaches the section, and ignores headings/answer cells.
  2. A reflow heuristic (always available): rebuild logical blocks from wrapped
     lines, classify headings vs questions, strip numbering and answer columns.
"""
import re
import json
from ..db import db, gen_id
from .. import config

# A line that starts a new item: "1.", "1.2", "1.2.3)", "(a)", "a)", "-", "•".
NUM_RE = re.compile(r"^\s*(?:\d+(?:\.\d+)*[.)]?|\([a-z0-9]{1,3}\)|[a-z][.)]|[-*•▪])\s+", re.I)
LEAD_RE = re.compile(r"^\s*(?:\d+(?:\.\d+)*[.)]?|\([a-z0-9]{1,3}\)|[a-z][.)]|[-*•▪])\s+", re.I)
# Trailing answer column / form artefacts to strip from a question.
ANSWER_TAIL = re.compile(
    r"\s*(?:[:\-–]\s*)?(?:y\s*/\s*n|yes\s*/\s*no|yes\s+no|true\s*/\s*false|n/?a|"
    r"☐|□|▢|◻|\[\s*\]|\(\s*\)|_{2,}|\.{3,}|select all that apply)\s*$", re.I)
INTERROGATIVE = re.compile(
    r"^(is|are|was|were|does|do|did|has|have|had|will|would|can|could|should|"
    r"may|might|please|provide|describe|list|state|specify|confirm|indicate|"
    r"name|identify|explain|detail|attach|upload|set out|give|"
    r"who|what|which|where|when|why|how|"
    # German (KYC forms are often German)
    r"ist|sind|hat|haben|liegt|liegen|besteht|bestehen|wird|werden|kann|"
    r"k[öo]nnen|gibt|bitte|geben|nennen|beschreiben|best[äa]tigen|erl[äa]utern|"
    r"f[üu]gen|legen|wer|was|wann|wo|wie|welche[rs]?)\b", re.I)


def infer_kind(prompt: str) -> str:
    p = prompt.lower()
    if re.search(r"\b(y/n|yes/no|yes or no|ja/nein)\b", p) or \
       re.match(r"^(is|are|does|do|has|have|will|ist|sind|hat|haben|liegt|liegen|besteht|gibt)\b", p):
        return "yesno"
    if re.search(r"\b(percent|percentage|%|shareholding|ownership stake|kapitalanteile|stimmanteile|beteiligung)\b", p):
        return "pct"
    if re.search(r"\b(date|incorporat|established|founded|datum|gegr[üu]ndet|gr[üu]ndung)\b", p):
        return "date"
    if re.search(r"\b(ubo|beneficial owner|controlling person|wirtschaftlich berechtigt)\w*\b", p):
        return "ubo_list"
    if re.search(r"\b(name of|legal name|company|entity|registered|firma|gesellschaft|bezeichnung)\b", p):
        return "entity"
    return "text"


def _is_heading(block: str) -> bool:
    """A section heading: short, no question mark, mostly upper-case (after any
    leading number is stripped) — e.g. "1. ENTITY & OWNERSHIP"."""
    core = LEAD_RE.sub("", block).strip().rstrip(":")
    if not core or core.endswith("?"):
        return False
    letters = [c for c in core if c.isalpha()]
    if not letters:
        return False
    upper = sum(c.isupper() for c in letters) / len(letters)
    if len(core) <= 70 and upper >= 0.6:
        return True
    # Title-cased short phrase with no verb (e.g. "Products & Services"). Skip
    # numbered lines — a wrapped question start like "1.2 Legal Entity Identifier"
    # is title-cased too, but it's content, not a section heading.
    if not LEAD_RE.match(block) and len(core) <= 45 and core[0].isupper() \
            and not INTERROGATIVE.match(core) and not core.endswith("?"):
        words = core.split()
        if 1 <= len(words) <= 5 and all(w[0].isupper() or not w[0].isalpha() for w in words):
            return True
    return False


def _looks_like_question(block: str) -> bool:
    core = LEAD_RE.sub("", block).strip()
    if len(core) < 8:
        return False
    if core.endswith("?"):
        return True
    return bool(INTERROGATIVE.match(core))


def _clean(block: str) -> str:
    s = LEAD_RE.sub("", block).strip()
    prev = None
    while prev != s:  # answer column can appear more than once after a join
        prev = s
        s = ANSWER_TAIL.sub("", s).strip()
    return re.sub(r"\s{2,}", " ", s).strip(" -–:")


def _reflow(raw: str):
    """Group physical lines into logical blocks: a new block begins at a blank
    line, a numbered/bulleted item, or a heading; wrapped lines join with a
    space."""
    blocks, cur = [], []

    def flush():
        if cur:
            blocks.append(" ".join(cur).strip())
            cur.clear()

    for line0 in raw.splitlines():
        line = line0.strip()
        if not line:
            flush()
            continue
        if _is_heading(line):
            flush()
            blocks.append(line)  # standalone heading block
            continue
        if NUM_RE.match(line) and cur:
            flush()
        cur.append(line)
    flush()
    return blocks


def parse_questions_heuristic(raw: str):
    out, section, seen = [], "", set()
    for block in _reflow(raw):
        if _is_heading(block):
            section = LEAD_RE.sub("", block).strip().rstrip(":")
            continue
        # A block may pack several questions / a trailing instruction; split on
        # every '?' boundary so "...policy? Please provide..." becomes two.
        candidates = [c.strip() for c in re.split(r"(?<=\?)\s+", block)] if "?" in block else [block]
        for cand in candidates:
            if not _looks_like_question(cand):
                continue
            prompt = _clean(cand)
            key = prompt.lower()
            if len(prompt) < 8 or key in seen:
                continue
            seen.add(key)
            out.append({"section": section, "prompt": prompt, "kind": infer_kind(prompt)})
    return out


def parse_questions_model(raw: str):
    """Model-assisted extraction (needs a key). Returns the same shape as the
    heuristic, or None if it can't run."""
    if not config.HAS_KEY:
        return None
    try:
        import anthropic
        sys = (
            "You extract questions from a KYC / due-diligence questionnaire. The text "
            "comes from a PDF/Word export, so questions may be split across lines, "
            "section headings sit between them, numbering may be separated, and answer "
            "cells like 'Y / N' trail the text. Reconstruct each ATOMIC question as one "
            "clean sentence. Ignore pure headings (use them as the 'section'), page "
            "numbers, and answer cells. Return STRICT JSON: "
            "{\"questions\":[{\"section\":str,\"prompt\":str}]}. Keep wording faithful; "
            "do not invent questions."
        )
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        resp = client.messages.create(model=config.MODEL, max_tokens=4000, system=sys,
                                       messages=[{"role": "user", "content": raw[:24000]}])
        txt = "".join(b.text for b in resp.content if getattr(b, "type", "") == "text")
        a, b = txt.find("{"), txt.rfind("}")
        data = json.loads(txt[a:b + 1])
        out, seen = [], set()
        for q in data.get("questions", []):
            prompt = re.sub(r"\s{2,}", " ", (q.get("prompt") or "").strip())
            if len(prompt) < 6 or prompt.lower() in seen:
                continue
            seen.add(prompt.lower())
            out.append({"section": (q.get("section") or "").strip(), "prompt": prompt,
                        "kind": infer_kind(prompt)})
        return out or None
    except Exception:
        return None


# a number cell such as "6", "6 a", "19 a1i" — anchors a sub-item's row
ANCHOR_RE = re.compile(r"^\d+(?:\s*[a-z]+\d*)*$", re.I)


def _cell_lines(page, bbox):
    """Text lines inside a cell bbox, each as (y, text), in reading order —
    wrapped words on the same visual line are joined."""
    x0, y0, x1, y1 = bbox[0], bbox[1], bbox[2], bbox[3]
    ws = [w for w in page.get_text("words")
          if w[0] >= x0 - 1 and w[2] <= x1 + 1 and w[1] >= y0 - 1 and w[3] <= y1 + 1]
    ws.sort(key=lambda w: (round(w[1]), w[0]))
    lines = []
    for w in ws:
        if lines and abs(w[1] - lines[-1][0]) < 4:
            lines[-1][1].append((w[0], w[4]))
        else:
            lines.append([w[1], [(w[0], w[4])]])
    return [(y, " ".join(t for _, t in sorted(toks))) for y, toks in lines]


def _slice_packed(page, no_bbox, q_bbox):
    """A cell that packs several numbered sub-items: align the Question column to
    the No column's anchor rows by vertical position, so wrapped question lines
    join under the right sub-item. Returns a list of question strings or None."""
    anchors = [(y, txt) for y, txt in _cell_lines(page, no_bbox) if ANCHOR_RE.match(txt.strip())]
    q_lines = _cell_lines(page, q_bbox)
    if len(anchors) < 2 or not q_lines:
        return None
    buckets = [[] for _ in anchors]
    for qy, qt in q_lines:
        # the sub-item a line belongs to is the anchor whose row it sits in —
        # the nearest anchor at or just above the line (half a row's tolerance)
        idx = 0
        for i, (ay, _) in enumerate(anchors):
            if qy >= ay - 7:
                idx = i
            else:
                break
        buckets[idx].append(qt)
    return [" ".join(b).strip() for b in buckets if " ".join(b).strip()]


def parse_questions_markdown(data: bytes):
    """Best path for tabular PDFs: pymupdf4llm renders the document as Markdown
    with full layout analysis, so each table row (No / Question / Answer) comes
    out clean — sub-items already on their own row and wrapped lines joined.
    Also captures the answer cell (an existing answer in the source). Returns the
    standard shape (with an extra 'answer'), or None."""
    try:
        import pymupdf4llm
        import fitz
    except Exception:
        return None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        md = pymupdf4llm.to_markdown(doc, show_progress=False)
    except Exception:
        return None

    def clean(s):
        return re.sub(r"\s{2,}", " ", (s or "").replace("<br>", " ").replace("**", "")).strip()

    out, section, seen = [], "", set()
    for line in md.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if not cells:
            continue
        # markdown separator row (|---|---|)
        if any("-" in c for c in cells) and all(set(c) <= set("-: ") for c in cells if c):
            continue
        no = clean(cells[0]) if len(cells) > 0 else ""
        q = clean(cells[1]) if len(cells) > 1 else ""
        ans = clean(cells[2]) if len(cells) > 2 else ""
        if no and not q:                              # a section-header row
            core = re.sub(r"^\d+[.)]\s*", "", no).strip()
            if re.match(r"^\d+[.)]\s", no) or _is_heading(no):
                section = core
            continue
        if not q or q.lower() in ("question", "no #", "answer"):
            continue
        prompt = clean(q)
        key = section + "|" + prompt.lower()
        if len(prompt) < 6 or key in seen:
            continue
        seen.add(key)
        out.append({"section": section, "prompt": prompt, "kind": infer_kind(prompt), "answer": ans})
    return out or None


def parse_questions_table(data: bytes):
    """Most KYC questionnaires are tables (No / Question / Answer). Reading the
    text flow scrambles those columns; pulling the table structure out of the
    PDF recovers clean questions and their sections. Returns the same shape, or
    None if no usable table is found."""
    try:
        import fitz
    except Exception:
        return None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception:
        return None
    out, section, seen = [], "", set()

    def add(sec, q):
        q = re.sub(r"\s{2,}", " ", q).strip(" -–")
        if len(q) < 6 or q.lower() == "question":
            return
        key = sec + "|" + q.lower()
        if key in seen:
            return
        seen.add(key)
        out.append({"section": sec, "prompt": q, "kind": infer_kind(q)})

    for pg in doc:
        try:
            tabs = pg.find_tables()
        except Exception:
            continue
        for t in tabs.tables:
            try:
                extracted = t.extract()
            except Exception:
                continue
            # locate the Question column (by header on the first page, else col 1)
            qcol = 1
            if t.header and t.header.names:
                names = [(n or "").strip().lower() for n in t.header.names]
                if "question" in names:
                    qcol = names.index("question")
            for row, r in zip(t.rows, extracted):
                no = (r[0] or "").strip() if len(r) > 0 else ""
                q = (r[qcol] or "").strip() if len(r) > qcol else ""
                if no and not q:                      # a section-header row
                    head = no.split("\n")[0].strip()  # ignore anything merged below it
                    core = re.sub(r"^\d+[.)]\s*", "", head).strip()
                    if _is_heading(head) or re.match(r"^\d+[.)]\s", head):
                        section = core
                    continue
                if not q:
                    continue
                no_lines = [x for x in no.split("\n") if x.strip()]
                cells = row.cells
                # a cell that packs several numbered sub-items → align by geometry
                if len(no_lines) > 1 and len(cells) > qcol and cells[0] and cells[qcol]:
                    try:
                        items = _slice_packed(pg, cells[0], cells[qcol])
                    except Exception:
                        items = None
                    if items:
                        for it in items:
                            add(section, it)
                        continue
                q_lines = [x.strip() for x in q.split("\n") if x.strip()]
                if len(no_lines) > 1 and len(no_lines) == len(q_lines):
                    for ql in q_lines:
                        add(section, ql)
                else:
                    add(section, " ".join(q_lines))
    return out or None


# Response-cell template lines ("Confirmed:", "Provided", "Yes", checkboxes) —
# scaffolding for the answer, not an answer.
TEMPLATE_LINE = re.compile(
    r"^\s*(?:[☐□▢◻✓xX]\s*)?(?:confirmed(?:\s+(?:yes|no))?|provided|not\s+(?:applicable|regulated)|n/?a|"
    r"yes|no|ja|nein|name|position|date|datum|title|signature|unterschrift|comments?)\s*[:.]?\s*$", re.I)
INSTRUCTION_LINE = re.compile(r"^\s*(?:if\b|please\b|falls\b|wenn\b|bitte\b|sofern\b)", re.I)


def _cell_answer(text: str) -> str:
    """Real content of a response cell — '' when it only holds templates
    ('Confirmed: / Provided:'), bare labels, option lists, checkboxes or
    conditional instructions."""
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    # an option list: several short label lines without values ("Hedge Fund" …)
    if len(lines) >= 3 and sum(1 for s in lines if len(s) <= 32 and ":" not in s) / len(lines) >= 0.6:
        return ""
    vals = []
    for s in lines:
        # side-by-side templates ("Provided     Not Regulated")
        segs = re.split(r"\s{2,}|\t", s)
        if all(TEMPLATE_LINE.match(x) for x in segs if x):
            continue
        if TEMPLATE_LINE.match(s) or INSTRUCTION_LINE.match(s) or re.match(r"^[^:]{2,60}:$", s):
            continue
        m = re.match(r"^([^:]{2,40}):\s*(.+)$", s)
        if m and TEMPLATE_LINE.match(m.group(1) + ":"):
            vals.append(m.group(2).strip())
            continue
        vals.append(s)
    return re.sub(r"\s{2,}", " ", " ".join(vals)).strip(" .")


def parse_questions_docx(content: bytes):
    """Word request lists (e.g. bank KYC requirements): section headings as
    paragraphs, then 2-column tables — left cell the request, right cell the
    response template (or a given answer). Walks body order so each table lands
    under its heading. Returns the standard shape or None."""
    try:
        from io import BytesIO
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn
        doc = Document(BytesIO(content))
    except Exception:
        return None
    out, section, seen = [], "", set()

    def add(prompt, answer=""):
        prompt = re.sub(r"\s{2,}", " ", prompt.replace("\n", " ")).strip(" -–")
        if len(prompt) < 6:
            return
        key = section + "|" + prompt.lower()
        if key in seen:
            return
        seen.add(key)
        out.append({"section": section, "prompt": prompt, "kind": infer_kind(prompt), "answer": answer})

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            t = Paragraph(child, doc).text.strip()
            if t and len(t) <= 70 and (t.endswith(":") or _is_heading(t)):
                section = t.rstrip(":").strip()
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                cells = row.cells
                if not cells:
                    continue
                left = cells[0].text.strip()
                if not left or left.lower() in ("question", "no", "no #", "item", "answer"):
                    continue
                # merged full-width row → sub-heading, not a request
                if len(cells) > 1 and cells[0]._tc is cells[-1]._tc:
                    if len(left) <= 70:
                        section = left.rstrip(":").strip()
                    continue
                right = cells[-1].text.strip() if len(cells) > 1 else ""
                add(left, _cell_answer(right))
    return out or None


def parse_questions_acroform(data: bytes):
    """Fillable PDF forms (AcroForm): the widgets themselves are the items to
    answer. Build each question from the text printed next to its field (same
    row to the left, else directly above, else directly below — German forms
    often label under the line); checkbox pairs (Ja/Nein) collapse onto the
    question line above them. Sections come from numbered headings (I. / 1.)."""
    try:
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception:
        return None
    widgets_total = sum(1 for pg in doc for _ in (pg.widgets() or []))
    if widgets_total < 3:
        return None
    # headings: roman numerals ("I. Identifizierung …") or upper/title-case lines,
    # NOT "2. Sind die Stimmanteile …" (numbered items are usually questions)
    HEAD = re.compile(r"^\s*[IVX]{1,4}\.\s+\S")
    GENERIC = re.compile(r"(?i)^(?:text|check\s*box|kontrollk[äa]stchen|feld|field|untitled|unbenannt)?[\s_\-]*\d*$")
    YESNO = re.compile(r"(?i)^\s*(?:ja|nein|yes|no)\b[\s:.]*(?:\(.*\))?\s*$")
    out, seen, section = [], set(), ""

    def add(prompt, kind):
        prompt = re.sub(r"[_\s]{2,}", " ", prompt).strip(" :;-–_")
        prompt = re.sub(r"Row\d+$", "", prompt).strip()      # form-field artefacts
        if len(prompt) < 4:
            return
        if prompt[0].islower() and len(prompt) < 40:          # wrapped-line fragment
            return
        if re.match(r"^\d+\s", prompt) and len(prompt) < 40:  # footnote ("1 Vgl. …")
            return
        key = prompt.lower()
        # near-duplicates from checkbox pairs: one label often contains the other
        if any(key in k or k in key for k in seen):
            return
        seen.add(key)
        out.append({"section": section, "prompt": prompt, "kind": kind, "answer": ""})

    for page in doc:
        widgets = sorted((page.widgets() or []), key=lambda w: (round(w.rect.y0), w.rect.x0))
        if not widgets:
            continue
        words = page.get_text("words")
        # words grouped into visual lines
        lines = []
        for w in sorted(words, key=lambda w: (round(w[1] / 4), w[0])):
            if lines and abs(w[1] - lines[-1][0]) < 4:
                lines[-1][2].append(w)
            else:
                lines.append([w[1], w[3], [w]])
        vlines = [(y0, y1, " ".join(t[4] for t in toks)) for y0, y1, toks in lines]

        def nearby(r):
            same = [t[4] for t in words if t[2] <= r.x0 + 2 and r.x0 - t[2] < 260
                    and t[1] < r.y1 - 1 and t[3] > r.y0 + 1]
            if same:
                return " ".join(same)
            # up to two wrapped label lines directly above, in reading order
            above = sorted((t for t in words if t[3] <= r.y0 + 2 and r.y0 - t[3] < 30
                            and t[0] > r.x0 - 200 and t[2] < r.x1 + 200),
                           key=lambda t: (round(t[1] / 4), t[0]))
            if above:
                return " ".join(t[4] for t in above)
            below = [t[4] for t in words if t[1] >= r.y1 - 2 and t[1] - r.y1 < 14
                     and t[0] > r.x0 - 180 and t[2] < r.x1 + 180]
            return " ".join(below)

        for w in widgets:
            r = w.rect
            # running section: nearest heading line above the widget
            heads = [t for (y0, y1, t) in vlines if y1 <= r.y0 and HEAD.match(t) and len(t) <= 90]
            if heads:
                section = re.sub(r"^\s*(?:[IVX]{1,4}\.|\d{1,2}\.)\s*", "", heads[-1]).strip()
            label = nearby(r)
            name = (w.field_name or "").strip()
            if not label and not GENERIC.match(name):
                label = name.replace("_", " ")
            import fitz as _f
            is_box = w.field_type == _f.PDF_WIDGET_TYPE_CHECKBOX
            if is_box and YESNO.match(label or ""):
                # Ja/Nein pair → the question is the text just above (join up to
                # two wrapped lines so the full sentence survives)
                cand = [(y0, t) for (y0, y1, t) in vlines if y1 <= r.y0 + 2 and len(t) > 12
                        and not YESNO.match(t)]
                if cand:
                    qlines = [t for _, t in cand[-2:]]
                    if len(qlines) == 2 and cand[-1][0] - cand[-2][0] > 30:
                        qlines = qlines[-1:]          # not adjacent → single line
                    add(" ".join(qlines), "yesno")
                continue
            if label:
                add(label, "yesno" if is_box else infer_kind(label))
    return out or None


def docx_qa_pairs(content: bytes):
    """Extract (question, answer) pairs from an already-answered Word
    questionnaire — Question/Answer tables and question→answer paragraph pairs —
    for feeding the KYC Brain. Returns a list of (prompt, answer)."""
    from io import BytesIO
    from docx import Document
    try:
        doc = Document(BytesIO(content))
    except Exception:
        return []
    pairs = []
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] and cells[0].lower() not in ("question", "no", "no #"):
                ans = next((c for c in cells[1:] if c), "")
                pairs.append((cells[0], ans))
    paras = [p.text.strip() for p in doc.paragraphs]
    for i, t in enumerate(paras):
        if t and (t.endswith("?") or _looks_like_question(t)) and i + 1 < len(paras):
            nxt = paras[i + 1]
            if nxt and not nxt.endswith("?"):
                pairs.append((t, nxt))
    return pairs


def parse_document(raw_text: str, filename: str = "", content: bytes = b""):
    """Pick the best extractor: table structure for PDFs, then model, then the
    text reflow heuristic."""
    name = (filename or "").lower()
    if name.endswith(".docx") and content:
        d = parse_questions_docx(content)             # request lists in 2-col tables
        if d and len(d) >= 3:
            return d
    if name.endswith(".pdf") and content:
        af = parse_questions_acroform(content)        # fillable form → fields are the items
        if af and len(af) >= 5:
            return af
        md = parse_questions_markdown(content)        # layout-aware, best for tables
        if md and len(md) >= 5:
            return md
        tabular = parse_questions_table(content)      # structural fallback
        if tabular and len(tabular) >= 5:
            return tabular
    return parse_questions(raw_text)


def parse_questions(raw: str):
    """Model first (if a key is configured), heuristic otherwise/as fallback."""
    return parse_questions_model(raw) or parse_questions_heuristic(raw)


def create_questionnaire(client_id: str, requester: str, title: str, raw_text: str,
                         filename: str = "", content: bytes = b"") -> str:
    qs = parse_document(raw_text, filename, content)
    qid = gen_id("qn")
    with db() as con:
        con.execute("INSERT INTO questionnaires (id, client_id, requester, title, status) VALUES (?,?,?,?,'parsed')",
                    (qid, client_id, requester, title))
        for i, q in enumerate(qs):
            con.execute("INSERT INTO questions (id, questionnaire_id, position, section, prompt, kind, source_answer) VALUES (?,?,?,?,?,?,?)",
                        (gen_id("q"), qid, i, q["section"], q["prompt"], q["kind"], q.get("answer", "")))
    return qid
